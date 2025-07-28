from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.shared import OxmlElement, qn
from fpdf import FPDF
from PIL import Image, ImageDraw
import os
import tempfile
from django.conf import settings
import base64
import subprocess
from pdf2image import convert_from_bytes
import warnings

warnings.filterwarnings("ignore", category=UserWarning, module="docx.styles.styles")

class ResumeGenerator:
    def __init__(self, resume_data):
        self.resume_data = self._clean_and_validate_data(resume_data)
        self.font_name = "Arial"
        self.heading_size = 14
        self.subheading_size = 12
        self.normal_size = 11
        self.profile_pic = self.resume_data.get('profile_pic')
        
        # Color schemes for different templates
        # Template 1 - Purple Modern (from first image)
        self.template1_colors = {
            'primary': RGBColor(89, 74, 158),      # Purple
            'secondary': RGBColor(102, 102, 102),   # Gray
            'accent': RGBColor(51, 72, 98),         # Dark blue
            'sidebar_bg': "E6E3F7"                  # Light purple
        }
        
        # Template 2 - Professional Blue (from second image)
        self.template2_colors = {
            'primary': RGBColor(51, 72, 98),        # Dark blue
            'secondary': RGBColor(102, 102, 102),   # Gray
            'accent': RGBColor(89, 74, 158),        # Purple
            'sidebar_bg': "334862",                 # Dark blue
            'light_bg': "F5F5F5"                   # Light gray
        }
        
        # Template 3 - Clean Minimalist (mixture)
        self.template3_colors = {
            'primary': RGBColor(51, 51, 51),        # Dark gray
            'secondary': RGBColor(102, 102, 102),   # Medium gray
            'accent': RGBColor(89, 74, 158),        # Purple accent
            'line_color': RGBColor(200, 200, 200)   # Light gray lines
        }

    def _clean_and_validate_data(self, data):
        """Clean and validate resume data, removing null/empty values and setting defaults"""
        cleaned_data = {}
        
        # Personal Information
        cleaned_data['full_name'] = self._safe_get(data, 'full_name', 'Your Name')
        cleaned_data['job_role'] = self._safe_get(data, 'job_role', 'Professional')
        cleaned_data['email'] = self._safe_get(data, 'email', 'your.email@example.com')
        cleaned_data['phone'] = self._safe_get(data, 'phone', '(555) 123-4567')
        cleaned_data['location'] = self._safe_get(data, 'location', 'Your City, State')
        cleaned_data['linkedin_url'] = self._safe_get(data, 'linkedin_url', '')
        cleaned_data['github_url'] = self._safe_get(data, 'github_url', '')
        cleaned_data['portfolio_url'] = self._safe_get(data, 'portfolio_url', '')
        
        # About Me / Summary
        cleaned_data['summary'] = self._safe_get(data, 'summary', 
            'Enthusiastic professional with a passion for excellence and continuous learning. '
            'Dedicated to delivering high-quality results and contributing to team success.')
        cleaned_data['profile_pic'] = data.get('profile_pic', '')
        
        # Skills - filter out empty skills
        skills = data.get('skills', [])
        cleaned_skills = []
        for skill in skills:
            if isinstance(skill, dict):
                skill_name = skill.get('name', '').strip()
                if skill_name and skill_name.lower() not in ['null', 'none', '']:
                    cleaned_skills.append({
                        'name': skill_name,
                        'level': skill.get('level', 'Intermediate')
                    })
            elif isinstance(skill, str) and skill.strip() and skill.lower() not in ['null', 'none']:
                cleaned_skills.append({'name': skill.strip(), 'level': 'Intermediate'})
        
        if not cleaned_skills:
            cleaned_skills = [
                {'name': 'Communication', 'level': 'Advanced'},
                {'name': 'Problem Solving', 'level': 'Advanced'},
                {'name': 'Teamwork', 'level': 'Advanced'},
                {'name': 'Time Management', 'level': 'Intermediate'},
                {'name': 'Leadership', 'level': 'Intermediate'}
            ]
        cleaned_data['skills'] = cleaned_skills
        
        # Work Experience - filter out empty experiences
        experiences = data.get('experiences', [])
        cleaned_experiences = []
        for exp in experiences:
            if isinstance(exp, dict):
                job_title = self._safe_get(exp, 'jobTitle', '') or self._safe_get(exp, 'job_title', '')
                company = self._safe_get(exp, 'company', '')
                
                if job_title.strip() and company.strip() and \
                   job_title.lower() not in ['null', 'none'] and company.lower() not in ['null', 'none']:
                    
                    responsibilities = exp.get('responsibilities', [])
                    cleaned_responsibilities = []
                    for resp in responsibilities:
                        if isinstance(resp, dict):
                            desc = resp.get('description', '').strip()
                        else:
                            desc = str(resp).strip()
                        
                        if desc and desc.lower() not in ['null', 'none', '']:
                            cleaned_responsibilities.append(desc)
                    
                    if not cleaned_responsibilities:
                        cleaned_responsibilities = [
                            f"Contributed to {company}'s success through dedicated work and collaboration.",
                            "Maintained high standards of professionalism and quality in all tasks.",
                            "Collaborated effectively with team members to achieve project goals."
                        ]
                    
                    cleaned_experiences.append({
                        'job_title': job_title,
                        'company': company,
                        'location': self._safe_get(exp, 'location', ''),
                        'duration': self._safe_get(exp, 'duration', 'Present'),
                        'responsibilities': cleaned_responsibilities
                    })
        
        if not cleaned_experiences:
            cleaned_experiences = [{
                'job_title': 'Recent Graduate / Job Seeker',
                'company': 'Currently Seeking Opportunities',
                'location': cleaned_data['location'],
                'duration': 'Present',
                'responsibilities': [
                    'Actively developing skills and seeking opportunities to contribute to a dynamic organization.',
                    'Engaged in continuous learning and professional development activities.',
                    'Building expertise through personal projects and online courses.'
                ]
            }]
        cleaned_data['experiences'] = cleaned_experiences
        
        # Education - filter out empty education
        educations = data.get('educations', [])
        cleaned_educations = []
        for edu in educations:
            if isinstance(edu, dict):
                institution = self._safe_get(edu, 'institution', '')
                degree = self._safe_get(edu, 'degree', '')
                
                if institution.strip() and degree.strip() and \
                   institution.lower() not in ['null', 'none'] and degree.lower() not in ['null', 'none']:
                    cleaned_educations.append({
                        'institution': institution,
                        'degree': degree,
                        'year': self._safe_get(edu, 'year', 'Year'),
                        'grade': self._safe_get(edu, 'grade', ''),
                        'course_name': self._safe_get(edu, 'course_name', degree)
                    })
        
        if not cleaned_educations:
            cleaned_educations = [{
                'institution': 'Educational Institution',
                'degree': 'Degree/Certification',
                'course_name': 'Relevant Course/Program',
                'year': 'Year',
                'grade': ''
            }]
        cleaned_data['educations'] = cleaned_educations
        
        # Projects - filter out empty projects
        projects = data.get('projects', [])
        cleaned_projects = []
        for proj in projects:
            if isinstance(proj, dict):
                title = self._safe_get(proj, 'title', '')
                description = self._safe_get(proj, 'description', '')
                
                if title.strip() and description.strip() and \
                   title.lower() not in ['null', 'none'] and description.lower() not in ['null', 'none']:
                    cleaned_projects.append({
                        'title': title,
                        'description': description,
                        'tech_stack': self._safe_get(proj, 'tech_stack', '') or self._safe_get(proj, 'techStack', ''),
                        'github_link': self._safe_get(proj, 'github_link', '') or self._safe_get(proj, 'githubLink', ''),
                        'summary': self._safe_get(proj, 'summary', description)
                    })
        
        if not cleaned_projects:
            cleaned_projects = [{
                'title': 'Personal Development Project',
                'description': 'Continuously working on personal and professional development through various learning initiatives and practical applications.',
                'tech_stack': 'Various Technologies and Tools',
                'github_link': '',
                'summary': 'Focused on skill enhancement and practical application of learned concepts.'
            }]
        cleaned_data['projects'] = cleaned_projects
        
        # Languages, Interests, and Additional Info
        cleaned_data['languages'] = self._clean_string_array(data.get('languages', []), ['English'])
        cleaned_data['interests'] = self._clean_string_array(data.get('interests', []), ['Learning', 'Technology', 'Innovation'])
        cleaned_data['hobbies'] = self._clean_string_array(data.get('hobbies', []), [])
        
        # Certifications and Achievements
        cleaned_data['certifications'] = self._clean_string_array(data.get('certifications', []), [])
        cleaned_data['achievements'] = self._clean_string_array(data.get('achievements', []), [])
        cleaned_data['volunteering'] = data.get('volunteering', [])
        
        return cleaned_data

    def _safe_get(self, data, key, default=''):
        """Safely get a value from dict, handling null/none/empty cases"""
        value = data.get(key, default)
        if not value or str(value).lower() in ['null', 'none', '']:
            return default
        return str(value).strip()

    def _clean_string_array(self, arr, default_items):
        """Clean array of strings, removing null/empty values"""
        cleaned = []
        for item in arr:
            if isinstance(item, dict):
                name = item.get('name', '').strip()
                if name and name.lower() not in ['null', 'none', '']:
                    cleaned.append(name)
            elif isinstance(item, str) and item.strip() and item.lower() not in ['null', 'none']:
                cleaned.append(item.strip())
        
        return cleaned if cleaned else default_items

    def _convert_to_pdf(self, docx_buffer, template_name):
        """Convert DOCX to PDF using LibreOffice or fallback to FPDF"""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_docx:
            temp_docx.write(docx_buffer.getvalue())
            temp_docx_path = temp_docx.name
        
        temp_pdf_path = temp_docx_path.replace('.docx', '.pdf')
        
        try:
            result = subprocess.run([
                r'C:\Program Files\LibreOffice\program\soffice.exe', '--headless', '--convert-to', 'pdf', 
                temp_docx_path, '--outdir', os.path.dirname(temp_docx_path)
            ], capture_output=True, timeout=30)
            
            if result.returncode == 0 and os.path.exists(temp_pdf_path):
                with open(temp_pdf_path, 'rb') as f:
                    pdf_buffer = BytesIO(f.read())
            else:
                raise Exception("LibreOffice conversion failed")
                
        except Exception as e:
            print(f"LibreOffice conversion failed: {e}. Using FPDF fallback.")
            pdf_buffer = self._create_fallback_pdf(template_name)
        
        try:
            os.unlink(temp_docx_path)
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
        except:
            pass
        
        return pdf_buffer

    def _create_fallback_pdf(self, template_name):
        """Create a simple PDF using FPDF as fallback"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font('Arial', 'B', 16)
        
        # Add name
        pdf.cell(0, 10, txt=self.resume_data.get('full_name', 'Resume'), ln=True, align='C')
        pdf.ln(5)
        
        # Contact info
        pdf.set_font('Arial', size=12)
        contact_info = [
            f"Email: {self.resume_data.get('email', '')}",
            f"Phone: {self.resume_data.get('phone', '')}",
            f"Location: {self.resume_data.get('location', '')}"
        ]
        
        for info in contact_info:
            if info.split(': ')[1]:
                pdf.cell(0, 8, txt=info, ln=True)
        pdf.ln(5)
        
        # Summary
        if self.resume_data.get('summary'):
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 8, txt="Professional Summary:", ln=True)
            pdf.set_font('Arial', size=10)
            
            summary = self.resume_data['summary']
            words = summary.split(' ')
            line = ""
            for word in words:
                test_line = line + word + " "
                if pdf.get_string_width(test_line) < 180:
                    line = test_line
                else:
                    if line:
                        pdf.cell(0, 6, txt=line.strip(), ln=True)
                    line = word + " "
            if line:
                pdf.cell(0, 6, txt=line.strip(), ln=True)
            pdf.ln(3)
        
        # Skills
        if self.resume_data.get('skills'):
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 8, txt="Skills:", ln=True)
            pdf.set_font('Arial', size=10)
            
            skills_text = ", ".join([skill['name'] for skill in self.resume_data['skills']])
            words = skills_text.split(' ')
            line = ""
            for word in words:
                test_line = line + word + " "
                if pdf.get_string_width(test_line) < 180:
                    line = test_line
                else:
                    if line:
                        pdf.cell(0, 6, txt=line.strip(), ln=True)
                    line = word + " "
            if line:
                pdf.cell(0, 6, txt=line.strip(), ln=True)
        
        buffer = BytesIO()
        pdf.output(buffer)
        buffer.seek(0)
        return buffer

    def _generate_thumbnail(self, pdf_buffer):
        """Generate thumbnail from PDF first page"""
        try:
            poppler_path = r'C:\poppler\poppler-24.08.0\Library\bin'
            
            pdf_buffer.seek(0)
            pages = convert_from_bytes(
                pdf_buffer.read(), 
                first_page=1, 
                last_page=1, 
                dpi=150,
                poppler_path=poppler_path
            )
            
            if pages:
                img = pages[0]
                img.thumbnail((300, 400), Image.Resampling.LANCZOS)
                
                thumbnail_buffer = BytesIO()
                img.save(thumbnail_buffer, format='PNG', optimize=True, quality=85)
                thumbnail_buffer.seek(0)
                return thumbnail_buffer
            
        except Exception as e:
            print(f"PDF thumbnail generation failed: {e}. Creating placeholder.")
        
        return self._create_placeholder_thumbnail()

    def _create_placeholder_thumbnail(self):
        """Create a placeholder thumbnail"""
        img = Image.new('RGB', (300, 400), color=(245, 245, 245))
        draw = ImageDraw.Draw(img)
        
        draw.rectangle([(10, 10), (290, 390)], outline=(200, 200, 200), width=2)
        
        try:
            from PIL import ImageFont
            font = ImageFont.load_default()
        except:
            font = None
        
        text = f"Resume\n{self.resume_data.get('full_name', 'Preview')}"
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        
        x = (300 - text_width) // 2
        y = (400 - text_height) // 2
        draw.text((x, y), text, fill=(100, 100, 100), font=font)
        
        thumbnail_buffer = BytesIO()
        img.save(thumbnail_buffer, format='PNG')
        thumbnail_buffer.seek(0)
        return thumbnail_buffer

    def generate_all_templates(self):
        """Generate all three templates and return data for database storage"""
        templates = []
        
        template_methods = [
            ('modern', self.generate_template1),
            ('professional', self.generate_template2),
            ('minimalist', self.generate_template3)
        ]
        
        for template_name, method in template_methods:
            try:
                print(f"Generating {template_name} template...")
                
                template_data = method()
                pdf_buffer = template_data['pdf']
                
                thumbnail_buffer = self._generate_thumbnail(pdf_buffer)
                
                pdf_buffer.seek(0)
                thumbnail_buffer.seek(0)
                
                template_info = {
                    'template_name': template_name,
                    'pdf_data': pdf_buffer.read(),
                    'pdf_filename': f"{template_name}_resume_{self.resume_data['full_name'].replace(' ', '_')}.pdf",
                    'thumbnail_data': thumbnail_buffer.read(),
                    'thumbnail_filename': f"{template_name}_thumbnail.png",
                }
                
                templates.append(template_info)
                print(f"✓ {template_name} template generated successfully")
                
            except Exception as e:
                print(f"✗ Error generating {template_name} template: {e}")
                continue
        
        return templates

    def generate_template1(self):
        """Template 1 - Modern Purple Design (Based on Image 1)"""
        doc = Document()
        
        # Set narrow margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # Create main layout table (2 columns: content | sidebar)
        main_table = doc.add_table(rows=1, cols=2)
        main_table.autofit = False
        main_table.columns[0].width = Inches(4.5)  # Main content
        main_table.columns[1].width = Inches(2.5)  # Sidebar
        
        self._remove_table_borders(main_table)
        
        left_cell = main_table.cell(0, 0)
        right_cell = main_table.cell(0, 1)
        
        # Set sidebar background color
        self._set_cell_background_color(right_cell, self.template1_colors['sidebar_bg'])
        
        # LEFT COLUMN (Main Content)
        self._add_name_and_role_template1(left_cell)
        self._add_about_me_template1(left_cell)
        self._add_work_experience_template1(left_cell)
        self._add_projects_template1(left_cell)
        
        # RIGHT COLUMN (Sidebar)
        self._add_personal_info_template1(right_cell)
        self._add_skills_template1(right_cell)
        self._add_education_template1(right_cell)
        self._add_languages_template1(right_cell)
        self._add_hobbies_template1(right_cell)
        
        return self._save_docx(doc, "modern")

    def generate_template2(self):
        """Template 2 - Professional with Profile Picture (Based on Image 2)"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)
            section.bottom_margin = Inches(0.4)
            section.left_margin = Inches(0.4)
            section.right_margin = Inches(0.4)
        
        # Create header with profile picture and basic info
        self._add_professional_header_template2(doc)
        
        # Create main content layout (sidebar | content)
        main_table = doc.add_table(rows=1, cols=2)
        main_table.autofit = False
        main_table.columns[0].width = Inches(2.5)  # Sidebar
        main_table.columns[1].width = Inches(4.5)  # Main content
        
        self._remove_table_borders(main_table)
        
        left_cell = main_table.cell(0, 0)
        right_cell = main_table.cell(0, 1)
        
        # Set sidebar background
        self._set_cell_background_color(left_cell, self.template2_colors['light_bg'])
        
        # LEFT COLUMN (Sidebar)
        self._add_skills_template2(left_cell)
        self._add_education_template2(left_cell)
        self._add_languages_template2(left_cell)
        self._add_certifications_template2(left_cell)
        
        # RIGHT COLUMN (Main Content)
        self._add_summary_template2(right_cell)
        self._add_experience_template2(right_cell)
        self._add_projects_template2(right_cell)
        self._add_achievements_template2(right_cell)
        
        return self._save_docx(doc, "professional")

    def generate_template3(self):
        """Template 3 - Clean Minimalist (Mixture of both)"""
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
        
        # Header with name and contact
        self._add_minimal_header_template3(doc)
        
        # Personal Information section
        self._add_personal_information_template3(doc)
        
        # About Me / Summary
        self._add_section_header_template3(doc, "ABOUT ME")
        summary_para = doc.add_paragraph(self.resume_data.get('summary', ''))
        summary_para.paragraph_format.space_after = Pt(16)
        
        # Work Experience
        self._add_section_header_template3(doc, "WORK EXPERIENCE")
        for exp in self.resume_data.get('experiences', []):
            self._add_experience_item_template3(doc, exp)
        
        # Education
        self._add_section_header_template3(doc, "EDUCATION")
        for edu in self.resume_data.get('educations', []):
            self._add_education_item_template3(doc, edu)
        
        # Projects
        self._add_section_header_template3(doc, "PROJECTS")
        for project in self.resume_data.get('projects', []):
            self._add_project_item_template3(doc, project)
        
        # Skills and Languages in columns
        skills_lang_table = doc.add_table(rows=1, cols=2)
        skills_lang_table.autofit = False
        skills_lang_table.columns[0].width = Inches(3.5)
        skills_lang_table.columns[1].width = Inches(3.5)
        self._remove_table_borders(skills_lang_table)
        
        skills_cell = skills_lang_table.cell(0, 0)
        lang_cell = skills_lang_table.cell(0, 1)
        
        self._add_skills_template3(skills_cell)
        self._add_languages_interests_template3(lang_cell)
        
        return self._save_docx(doc, "minimalist")

    def _save_docx(self, doc, template_name):
        """Save the document to a BytesIO object and convert to PDF"""
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        pdf_buffer = self._convert_to_pdf(buffer, template_name)
        return {
            'docx': buffer,
            'pdf': pdf_buffer,
            'template_name': template_name
        }

    # Helper methods for styling and layout
    def _remove_table_borders(self, table):
        """Remove all borders from a table"""
        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)

    def _set_cell_background_color(self, cell, color_hex):
        """Set background color for a table cell"""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color_hex)
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # TEMPLATE 1 METHODS - Modern Purple Design
    def _add_name_and_role_template1(self, cell):
        """Add name and job role for template 1"""
        # Name
        name_para = cell.add_paragraph()
        name_run = name_para.add_run(self.resume_data.get('full_name', ''))
        name_run.bold = True
        name_run.font.size = Pt(24)
        name_run.font.name = self.font_name
        name_run.font.color.rgb = self.template1_colors['primary']
        name_para.paragraph_format.space_after = Pt(0)
        
        # Job Role
        job_role = self.resume_data.get('job_role', 'Professional')
        if not job_role and self.resume_data.get('experiences'):
            job_role = self.resume_data['experiences'][0].get('job_title', 'Professional')
        
        role_para = cell.add_paragraph()
        role_run = role_para.add_run(job_role)
        role_run.font.size = Pt(14)
        role_run.font.name = self.font_name
        role_run.font.color.rgb = self.template1_colors['secondary']
        role_para.paragraph_format.space_after = Pt(16)

    def _add_about_me_template1(self, cell):
        """Add About Me section for template 1"""
        about_header = cell.add_paragraph()
        about_run = about_header.add_run("ABOUT ME")
        about_run.bold = True
        about_run.font.size = Pt(12)
        about_run.font.name = self.font_name
        about_run.font.color.rgb = self.template1_colors['primary']
        about_header.paragraph_format.space_before = Pt(8)
        about_header.paragraph_format.space_after = Pt(8)
        
        # Summary content
        summary_para = cell.add_paragraph()
        summary_run = summary_para.add_run(self.resume_data.get('summary', ''))
        summary_run.font.size = Pt(10)
        summary_run.font.name = self.font_name
        summary_para.paragraph_format.space_after = Pt(16)

    def _add_work_experience_template1(self, cell):
        """Add work experience section for template 1"""
        exp_header = cell.add_paragraph()
        exp_run = exp_header.add_run("WORK EXPERIENCE")
        exp_run.bold = True
        exp_run.font.size = Pt(12)
        exp_run.font.name = self.font_name
        exp_run.font.color.rgb = self.template1_colors['primary']
        exp_header.paragraph_format.space_before = Pt(16)
        exp_header.paragraph_format.space_after = Pt(8)
        
        for exp in self.resume_data.get('experiences', []):
            self._add_experience_item_template1(cell, exp)

    def _add_experience_item_template1(self, cell, exp):
        """Add individual experience item for template 1"""
        # Company and location
        company_para = cell.add_paragraph()
        company_text = f"{exp.get('company', '')}"
        if exp.get('location'):
            company_text += f", {exp.get('location', '')}"
        company_run = company_para.add_run(company_text)
        company_run.font.size = Pt(11)
        company_run.font.name = self.font_name
        company_para.paragraph_format.space_after = Pt(2)
        
        # Job title and duration
        title_para = cell.add_paragraph()
        title_run = title_para.add_run(exp.get('job_title', ''))
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.name = self.font_name
        
        if exp.get('duration'):
            duration_run = title_para.add_run(f"                    {exp.get('duration', '')}")
            duration_run.font.size = Pt(10)
            duration_run.font.name = self.font_name
            duration_run.font.color.rgb = self.template1_colors['secondary']
        
        title_para.paragraph_format.space_after = Pt(4)
        
        # Responsibilities
        for resp in exp.get('responsibilities', []):
            resp_para = cell.add_paragraph()
            resp_run = resp_para.add_run(f"• {resp}")
            resp_run.font.size = Pt(10)
            resp_run.font.name = self.font_name
            resp_para.paragraph_format.left_indent = Inches(0.2)
            resp_para.paragraph_format.space_after = Pt(2)
        
        cell.add_paragraph().paragraph_format.space_after = Pt(8)

    def _add_projects_template1(self, cell):
        """Add projects section for template 1"""
        proj_header = cell.add_paragraph()
        proj_run = proj_header.add_run("PROJECTS")
        proj_run.bold = True
        proj_run.font.size = Pt(12)
        proj_run.font.name = self.font_name
        proj_run.font.color.rgb = self.template1_colors['primary']
        proj_header.paragraph_format.space_before = Pt(16)
        proj_header.paragraph_format.space_after = Pt(8)
        
        for project in self.resume_data.get('projects', []):
            self._add_project_item_template1(cell, project)

    def _add_project_item_template1(self, cell, project):
        """Add individual project item for template 1"""
        # Project name
        proj_para = cell.add_paragraph()
        proj_run = proj_para.add_run(project.get('title', ''))
        proj_run.bold = True
        proj_run.font.size = Pt(11)
        proj_run.font.name = self.font_name
        proj_para.paragraph_format.space_after = Pt(4)
        
        # Tech stack
        if project.get('tech_stack'):
            tech_para = cell.add_paragraph()
            tech_run = tech_para.add_run(f"Tech Stack: {project.get('tech_stack', '')}")
            tech_run.font.size = Pt(9)
            tech_run.font.name = self.font_name
            tech_run.font.color.rgb = self.template1_colors['secondary']
            tech_para.paragraph_format.space_after = Pt(2)
        
        # Project summary/description
        desc_para = cell.add_paragraph()
        desc_run = desc_para.add_run(f"• {project.get('summary', '') or project.get('description', '')}")
        desc_run.font.size = Pt(10)
        desc_run.font.name = self.font_name
        desc_para.paragraph_format.left_indent = Inches(0.2)
        desc_para.paragraph_format.space_after = Pt(8)

    def _add_personal_info_template1(self, cell):
        """Add personal information section to sidebar for template 1"""
        contact_header = cell.add_paragraph()
        contact_run = contact_header.add_run("PERSONAL INFORMATION")
        contact_run.bold = True
        contact_run.font.size = Pt(11)
        contact_run.font.name = self.font_name
        contact_run.font.color.rgb = self.template1_colors['primary']
        contact_header.paragraph_format.space_before = Pt(8)
        contact_header.paragraph_format.space_after = Pt(8)
        
        # Contact details
        contact_items = []
        if self.resume_data.get('phone'):
            contact_items.append(f"📞 {self.resume_data.get('phone')}")
        if self.resume_data.get('email'):
            contact_items.append(f"✉ {self.resume_data.get('email')}")
        if self.resume_data.get('location'):
            contact_items.append(f"📍 {self.resume_data.get('location')}")
        if self.resume_data.get('linkedin_url'):
            contact_items.append("🔗 LinkedIn Profile")
        if self.resume_data.get('github_url'):
            contact_items.append("💻 GitHub Profile")
        if self.resume_data.get('portfolio_url'):
            contact_items.append("🌐 Portfolio Link")
        
        for item in contact_items:
            contact_para = cell.add_paragraph()
            contact_item_run = contact_para.add_run(item)
            contact_item_run.font.size = Pt(9)
            contact_item_run.font.name = self.font_name
            contact_para.paragraph_format.space_after = Pt(3)

    def _add_skills_template1(self, cell):
        """Add skills section to sidebar for template 1"""
        skills_header = cell.add_paragraph()
        skills_run = skills_header.add_run("SKILLS")
        skills_run.bold = True
        skills_run.font.size = Pt(11)
        skills_run.font.name = self.font_name
        skills_run.font.color.rgb = self.template1_colors['primary']
        skills_header.paragraph_format.space_before = Pt(16)
        skills_header.paragraph_format.space_after = Pt(8)
        
        for skill in self.resume_data.get('skills', []):
            skill_para = cell.add_paragraph()
            skill_text = f"• {skill.get('name', '')}"
            if skill.get('level') and skill.get('level') != 'Intermediate':
                skill_text += f" ({skill.get('level', '')})"
            skill_run = skill_para.add_run(skill_text)
            skill_run.font.size = Pt(9)
            skill_run.font.name = self.font_name
            skill_para.paragraph_format.space_after = Pt(3)

    def _add_education_template1(self, cell):
        """Add education section to sidebar for template 1"""
        edu_header = cell.add_paragraph()
        edu_run = edu_header.add_run("EDUCATION")
        edu_run.bold = True
        edu_run.font.size = Pt(11)
        edu_run.font.name = self.font_name
        edu_run.font.color.rgb = self.template1_colors['primary']
        edu_header.paragraph_format.space_before = Pt(16)
        edu_header.paragraph_format.space_after = Pt(8)
        
        for edu in self.resume_data.get('educations', []):
            # Course/Degree name
            course_para = cell.add_paragraph()
            course_run = course_para.add_run(edu.get('course_name', '') or edu.get('degree', ''))
            course_run.bold = True
            course_run.font.size = Pt(10)
            course_run.font.name = self.font_name
            course_para.paragraph_format.space_after = Pt(2)
            
            # Institution
            inst_para = cell.add_paragraph()
            inst_run = inst_para.add_run(edu.get('institution', ''))
            inst_run.font.size = Pt(9)
            inst_run.font.name = self.font_name
            inst_para.paragraph_format.space_after = Pt(2)
            
            # Year and grade
            details = []
            if edu.get('year'):
                details.append(edu['year'])
            if edu.get('grade'):
                details.append(f"Grade: {edu['grade']}")
            
            if details:
                details_para = cell.add_paragraph()
                details_run = details_para.add_run(" — ".join(details))
                details_run.font.size = Pt(9)
                details_run.font.name = self.font_name
                details_run.font.color.rgb = self.template1_colors['secondary']
                details_para.paragraph_format.space_after = Pt(8)

    def _add_languages_template1(self, cell):
        """Add languages section to sidebar for template 1"""
        if not self.resume_data.get('languages'):
            return
            
        lang_header = cell.add_paragraph()
        lang_run = lang_header.add_run("LANGUAGES")
        lang_run.bold = True
        lang_run.font.size = Pt(11)
        lang_run.font.name = self.font_name
        lang_run.font.color.rgb = self.template1_colors['primary']
        lang_header.paragraph_format.space_before = Pt(16)
        lang_header.paragraph_format.space_after = Pt(8)
        
        for lang in self.resume_data.get('languages', []):
            lang_para = cell.add_paragraph()
            lang_run = lang_para.add_run(f"• {lang}")
            lang_run.font.size = Pt(9)
            lang_run.font.name = self.font_name
            lang_para.paragraph_format.space_after = Pt(3)

    def _add_hobbies_template1(self, cell):
        """Add hobbies/interests section to sidebar for template 1"""
        interests = self.resume_data.get('interests', []) + self.resume_data.get('hobbies', [])
        if not interests:
            return
            
        hobbies_header = cell.add_paragraph()
        hobbies_run = hobbies_header.add_run("HOBBIES & INTERESTS")
        hobbies_run.bold = True
        hobbies_run.font.size = Pt(11)
        hobbies_run.font.name = self.font_name
        hobbies_run.font.color.rgb = self.template1_colors['primary']
        hobbies_header.paragraph_format.space_before = Pt(16)
        hobbies_header.paragraph_format.space_after = Pt(8)
        
        for interest in interests:
            int_para = cell.add_paragraph()
            int_run = int_para.add_run(f"• {interest}")
            int_run.font.size = Pt(9)
            int_run.font.name = self.font_name
            int_para.paragraph_format.space_after = Pt(3)

    # TEMPLATE 2 METHODS - Professional with Profile Picture
    def _add_professional_header_template2(self, doc):
        """Add professional header with profile picture and name for template 2"""
        # Create header table for layout
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.5)  # Profile picture
        header_table.columns[1].width = Inches(5.5)  # Name and contact info
        
        self._remove_table_borders(header_table)
        
        pic_cell = header_table.cell(0, 0)
        info_cell = header_table.cell(0, 1)
        
        # Set profile picture cell background
        self._set_cell_background_color(pic_cell, self.template2_colors['sidebar_bg'])
        
        # Add profile picture or placeholder
        if self.profile_pic:
            self._add_circular_image(pic_cell, self.profile_pic, size=1.2)
        else:
            pic_para = pic_cell.add_paragraph()
            pic_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            pic_para.paragraph_format.space_before = Pt(15)
            pic_para.paragraph_format.space_after = Pt(15)
            placeholder_run = pic_para.add_run("Profile\nPicture")
            placeholder_run.font.size = Pt(9)
            placeholder_run.font.name = self.font_name
            placeholder_run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        
        # Name and job role in info cell
        name_para = info_cell.add_paragraph()
        name_run = name_para.add_run(self.resume_data.get('full_name', ''))
        name_run.bold = True
        name_run.font.size = Pt(20)
        name_run.font.name = self.font_name
        name_run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        name_para.paragraph_format.space_after = Pt(2)
        
        # Job role
        job_role = self.resume_data.get('job_role', 'Professional')
        if not job_role and self.resume_data.get('experiences'):
            job_role = self.resume_data['experiences'][0].get('job_title', 'Professional')
        
        role_para = info_cell.add_paragraph()
        role_run = role_para.add_run(job_role.upper())
        role_run.font.size = Pt(12)
        role_run.font.name = self.font_name
        role_run.font.color.rgb = self.template2_colors['secondary']
        role_para.paragraph_format.space_after = Pt(8)
        
        # Basic contact info
        contact_info = []
        if self.resume_data.get('phone'):
            contact_info.append(f"📞 {self.resume_data.get('phone')}")
        if self.resume_data.get('email'):
            contact_info.append(f"✉ {self.resume_data.get('email')}")
        
        for contact in contact_info:
            contact_para = info_cell.add_paragraph()
            contact_run = contact_para.add_run(contact)
            contact_run.font.size = Pt(10)
            contact_run.font.name = self.font_name
            contact_para.paragraph_format.space_after = Pt(2)

    def _add_summary_template2(self, cell):
        """Add summary section for template 2"""
        summary_header = cell.add_paragraph()
        summary_run = summary_header.add_run("SUMMARY")
        summary_run.bold = True
        summary_run.font.size = Pt(12)
        summary_run.font.name = self.font_name
        summary_run.font.color.rgb = self.template2_colors['primary']
        summary_header.paragraph_format.space_before = Pt(8)
        summary_header.paragraph_format.space_after = Pt(6)
        
        summary_para = cell.add_paragraph(self.resume_data.get('summary', ''))
        summary_para.paragraph_format.space_after = Pt(12)

    def _add_experience_template2(self, cell):
        """Add experience section for template 2"""
        exp_header = cell.add_paragraph()
        exp_run = exp_header.add_run("EXPERIENCE")
        exp_run.bold = True
        exp_run.font.size = Pt(12)
        exp_run.font.name = self.font_name
        exp_run.font.color.rgb = self.template2_colors['primary']
        exp_header.paragraph_format.space_before = Pt(12)
        exp_header.paragraph_format.space_after = Pt(6)
        
        for exp in self.resume_data.get('experiences', []):
            self._add_experience_item_template2(cell, exp)

    def _add_experience_item_template2(self, cell, exp):
        """Add individual experience item for template 2"""
        # Company name and year
        company_para = cell.add_paragraph()
        company_text = f"{exp.get('company', '')}"
        if exp.get('duration'):
            company_text += f"                    {exp.get('duration', '')}"
        company_run = company_para.add_run(company_text)
        company_run.font.size = Pt(11)
        company_run.font.name = self.font_name
        company_para.paragraph_format.space_after = Pt(2)
        
        # Job role and location
        role_para = cell.add_paragraph()
        role_text = exp.get('job_title', '')
        if exp.get('location'):
            role_text += f", {exp.get('location', '')}"
        role_run = role_para.add_run(role_text)
        role_run.bold = True
        role_run.font.size = Pt(11)
        role_run.font.name = self.font_name
        role_para.paragraph_format.space_after = Pt(4)
        
        # Summary/Responsibilities
        for resp in exp.get('responsibilities', []):
            resp_para = cell.add_paragraph()
            resp_run = resp_para.add_run(f"• {resp}")
            resp_run.font.size = Pt(10)
            resp_run.font.name = self.font_name
            resp_para.paragraph_format.left_indent = Inches(0.2)
            resp_para.paragraph_format.space_after = Pt(2)
        
        cell.add_paragraph().paragraph_format.space_after = Pt(8)

    def _add_projects_template2(self, cell):
        """Add projects section for template 2"""
        proj_header = cell.add_paragraph()
        proj_run = proj_header.add_run("PROJECTS")
        proj_run.bold = True
        proj_run.font.size = Pt(12)
        proj_run.font.name = self.font_name
        proj_run.font.color.rgb = self.template2_colors['primary']
        proj_header.paragraph_format.space_before = Pt(12)
        proj_header.paragraph_format.space_after = Pt(6)
        
        for project in self.resume_data.get('projects', []):
            self._add_project_item_template2(cell, project)

    def _add_project_item_template2(self, cell, project):
        """Add individual project item for template 2"""
        # Project name with link indicator
        proj_para = cell.add_paragraph()
        proj_text = project.get('title', '')
        if project.get('github_link'):
            proj_text += "                    Link"
        proj_run = proj_para.add_run(proj_text)
        proj_run.bold = True
        proj_run.font.size = Pt(11)
        proj_run.font.name = self.font_name
        proj_para.paragraph_format.space_after = Pt(2)
        
        # Tech stack
        if project.get('tech_stack'):
            tech_para = cell.add_paragraph()
            tech_run = tech_para.add_run(f"Tech Stack: {project.get('tech_stack', '')}")
            tech_run.font.size = Pt(9)
            tech_run.font.name = self.font_name
            tech_run.font.color.rgb = self.template2_colors['secondary']
            tech_para.paragraph_format.space_after = Pt(2)
        
        # Project summary
        desc_para = cell.add_paragraph()
        desc_run = desc_para.add_run(f"• {project.get('summary', '') or project.get('description', '')}")
        desc_run.font.size = Pt(10)
        desc_run.font.name = self.font_name
        desc_para.paragraph_format.left_indent = Inches(0.2)
        desc_para.paragraph_format.space_after = Pt(8)

    def _add_achievements_template2(self, cell):
        """Add achievements section for template 2"""
        achievements = self.resume_data.get('achievements', [])
        if not achievements:
            return
            
        ach_header = cell.add_paragraph()
        ach_run = ach_header.add_run("ACHIEVEMENTS")
        ach_run.bold = True
        ach_run.font.size = Pt(12)
        ach_run.font.name = self.font_name
        ach_run.font.color.rgb = self.template2_colors['primary']
        ach_header.paragraph_format.space_before = Pt(12)
        ach_header.paragraph_format.space_after = Pt(6)
        
        for achievement in achievements:
            ach_para = cell.add_paragraph()
            ach_run = ach_para.add_run(f"• {achievement}")
            ach_run.font.size = Pt(10)
            ach_run.font.name = self.font_name
            ach_para.paragraph_format.left_indent = Inches(0.2)
            ach_para.paragraph_format.space_after = Pt(2)

    def _add_skills_template2(self, cell):
        """Add skills section for template 2"""
        skills_header = cell.add_paragraph()
        skills_run = skills_header.add_run("SKILLS")
        skills_run.bold = True
        skills_run.font.size = Pt(11)
        skills_run.font.name = self.font_name
        skills_run.font.color.rgb = self.template2_colors['primary']
        skills_header.paragraph_format.space_before = Pt(8)
        skills_header.paragraph_format.space_after = Pt(8)
        
        for skill in self.resume_data.get('skills', []):
            skill_para = cell.add_paragraph()
            # Add skill with level indicator (using progress bars like in image)
            skill_text = f"• {skill.get('name', '')}"
            level = skill.get('level', 'Intermediate')
            if level:
                if level.lower() == 'advanced':
                    skill_text += " ●●●●●"
                elif level.lower() == 'intermediate':
                    skill_text += " ●●●○○"
                else:  # Beginner
                    skill_text += " ●●○○○"
            
            skill_run = skill_para.add_run(skill_text)
            skill_run.font.size = Pt(9)
            skill_run.font.name = self.font_name
            skill_para.paragraph_format.space_after = Pt(4)

    def _add_education_template2(self, cell):
        """Add education section for template 2"""
        edu_header = cell.add_paragraph()
        edu_run = edu_header.add_run("EDUCATION")
        edu_run.bold = True
        edu_run.font.size = Pt(11)
        edu_run.font.name = self.font_name
        edu_run.font.color.rgb = self.template2_colors['primary']
        edu_header.paragraph_format.space_before = Pt(16)
        edu_header.paragraph_format.space_after = Pt(8)
        
        for edu in self.resume_data.get('educations', []):
            # Course name
            course_para = cell.add_paragraph()
            course_run = course_para.add_run(edu.get('course_name', '') or edu.get('degree', ''))
            course_run.bold = True
            course_run.font.size = Pt(10)
            course_run.font.name = self.font_name
            course_para.paragraph_format.space_after = Pt(2)
            
            # Institution name
            inst_para = cell.add_paragraph()
            inst_run = inst_para.add_run(edu.get('institution', ''))
            inst_run.font.size = Pt(9)
            inst_run.font.name = self.font_name
            inst_para.paragraph_format.space_after = Pt(2)
            
            # Year and score
            details = []
            if edu.get('year'):
                details.append(edu['year'])
            if edu.get('grade'):
                details.append(f"Score: {edu['grade']}")
            
            if details:
                details_para = cell.add_paragraph()
                details_run = details_para.add_run(" — ".join(details))
                details_run.font.size = Pt(9)
                details_run.font.name = self.font_name
                details_run.font.color.rgb = self.template2_colors['secondary']
                details_para.paragraph_format.space_after = Pt(8)

    def _add_languages_template2(self, cell):
        """Add languages section for template 2"""
        if not self.resume_data.get('languages'):
            return
            
        lang_header = cell.add_paragraph()
        lang_run = lang_header.add_run("LANGUAGES")
        lang_run.bold = True
        lang_run.font.size = Pt(11)
        lang_run.font.name = self.font_name
        lang_run.font.color.rgb = self.template2_colors['primary']
        lang_header.paragraph_format.space_before = Pt(16)
        lang_header.paragraph_format.space_after = Pt(8)
        
        for lang in self.resume_data.get('languages', []):
            lang_para = cell.add_paragraph()
            lang_run = lang_para.add_run(f"• {lang}")
            lang_run.font.size = Pt(9)
            lang_run.font.name = self.font_name
            lang_para.paragraph_format.space_after = Pt(3)

    def _add_certifications_template2(self, cell):
        """Add certifications section for template 2"""
        certs = self.resume_data.get('certifications', [])
        if not certs:
            return
            
        cert_header = cell.add_paragraph()
        cert_run = cert_header.add_run("CERTIFICATIONS")
        cert_run.bold = True
        cert_run.font.size = Pt(11)
        cert_run.font.name = self.font_name
        cert_run.font.color.rgb = self.template2_colors['primary']
        cert_header.paragraph_format.space_before = Pt(16)
        cert_header.paragraph_format.space_after = Pt(8)
        
        for cert in certs:
            cert_para = cell.add_paragraph()
            cert_text = cert if isinstance(cert, str) else cert.get('name', '')
            cert_run = cert_para.add_run(f"• {cert_text}")
            cert_run.font.size = Pt(9)
            cert_run.font.name = self.font_name
            cert_para.paragraph_format.space_after = Pt(3)

    # TEMPLATE 3 METHODS - Clean Minimalist
    def _add_minimal_header_template3(self, doc):
        """Add minimal header with name and job role for template 3"""
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(self.resume_data.get('full_name', ''))
        name_run.bold = True
        name_run.font.size = Pt(22)
        name_run.font.name = self.font_name
        name_run.font.color.rgb = self.template3_colors['primary']
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        name_para.paragraph_format.space_after = Pt(2)
        
        # Job Role
        job_role = self.resume_data.get('job_role', 'Professional')
        if not job_role and self.resume_data.get('experiences'):
            job_role = self.resume_data['experiences'][0].get('job_title', 'Professional')
        
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(job_role)
        role_run.font.size = Pt(12)
        role_run.font.name = self.font_name
        role_run.font.color.rgb = self.template3_colors['secondary']
        role_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        role_para.paragraph_format.space_after = Pt(12)

    def _add_personal_information_template3(self, doc):
        """Add personal information section for template 3"""
        # Contact info in a clean line
        contact_info = []
        if self.resume_data.get('phone'):
            contact_info.append(self.resume_data['phone'])
        if self.resume_data.get('email'):
            contact_info.append(self.resume_data['email'])
        if self.resume_data.get('location'):
            contact_info.append(self.resume_data['location'])
        if self.resume_data.get('linkedin_url'):
            contact_info.append("LinkedIn")
        if self.resume_data.get('github_url'):
            contact_info.append("GitHub")
        
        if contact_info:
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(" | ".join(contact_info))
            contact_run.font.size = Pt(10)
            contact_run.font.name = self.font_name
            contact_run.font.color.rgb = self.template3_colors['secondary']
            contact_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            contact_para.paragraph_format.space_after = Pt(16)

    def _add_section_header_template3(self, doc, title):
        """Add section header for template 3"""
        header_para = doc.add_paragraph()
        header_run = header_para.add_run(title)
        header_run.bold = True
        header_run.font.size = Pt(12)
        header_run.font.name = self.font_name
        header_run.font.color.rgb = self.template3_colors['accent']
        header_para.paragraph_format.space_before = Pt(16)
        header_para.paragraph_format.space_after = Pt(6)
        
        # Add simple underline
        underline_para = doc.add_paragraph()
        underline_run = underline_para.add_run("_" * 70)
        underline_run.font.size = Pt(8)
        underline_run.font.color.rgb = self.template3_colors['line_color']
        underline_para.paragraph_format.space_after = Pt(8)

    def _add_experience_item_template3(self, doc, exp):
        """Add experience item for template 3"""
        # Job title and company - Year
        title_para = doc.add_paragraph()
        title_text = f"{exp.get('job_title', '')} — {exp.get('company', '')}"
        if exp.get('duration'):
            title_text += f"                    {exp.get('duration', '')}"
        title_run = title_para.add_run(title_text)
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.name = self.font_name
        title_para.paragraph_format.space_after = Pt(2)
        
        # Location if available
        if exp.get('location'):
            loc_para = doc.add_paragraph()
            loc_run = loc_para.add_run(exp.get('location', ''))
            loc_run.font.size = Pt(10)
            loc_run.font.name = self.font_name
            loc_run.font.color.rgb = self.template3_colors['secondary']
            loc_para.paragraph_format.space_after = Pt(4)
        
        # Responsibilities
        for resp in exp.get('responsibilities', []):
            resp_para = doc.add_paragraph()
            resp_run = resp_para.add_run(f"• {resp}")
            resp_run.font.size = Pt(10)
            resp_run.font.name = self.font_name
            resp_para.paragraph_format.left_indent = Inches(0.25)
            resp_para.paragraph_format.space_after = Pt(2)
        
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _add_education_item_template3(self, doc, edu):
        """Add education item for template 3"""
        # Course/Degree name and Institution - Year
        edu_para = doc.add_paragraph()
        edu_text = f"{edu.get('course_name', '') or edu.get('degree', '')} — {edu.get('institution', '')}"
        if edu.get('year'):
            edu_text += f"                    {edu.get('year', '')}"
        edu_run = edu_para.add_run(edu_text)
        edu_run.bold = True
        edu_run.font.size = Pt(11)
        edu_run.font.name = self.font_name
        edu_para.paragraph_format.space_after = Pt(2)
        
        # Grade if available
        if edu.get('grade'):
            grade_para = doc.add_paragraph()
            grade_run = grade_para.add_run(f"Grade: {edu.get('grade', '')}")
            grade_run.font.size = Pt(10)
            grade_run.font.name = self.font_name
            grade_run.font.color.rgb = self.template3_colors['secondary']
            grade_para.paragraph_format.space_after = Pt(8)
        else:
            doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _add_project_item_template3(self, doc, project):
        """Add project item for template 3"""
        # Project name with link indicator
        proj_para = doc.add_paragraph()
        proj_text = project.get('title', '')
        if project.get('github_link'):
            proj_text += "                    Link"
        proj_run = proj_para.add_run(proj_text)
        proj_run.bold = True
        proj_run.font.size = Pt(11)
        proj_run.font.name = self.font_name
        proj_para.paragraph_format.space_after = Pt(2)
        
        # Tech stack
        if project.get('tech_stack'):
            tech_para = doc.add_paragraph()
            tech_run = tech_para.add_run(f"Tech Stack: {project.get('tech_stack', '')}")
            tech_run.font.size = Pt(9)
            tech_run.font.name = self.font_name
            tech_run.font.color.rgb = self.template3_colors['secondary']
            tech_para.paragraph_format.space_after = Pt(2)
        
        # Project summary
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(f"• {project.get('summary', '') or project.get('description', '')}")
        summary_run.font.size = Pt(10)
        summary_run.font.name = self.font_name
        summary_para.paragraph_format.left_indent = Inches(0.25)
        summary_para.paragraph_format.space_after = Pt(8)

    def _add_skills_template3(self, cell):
        """Add skills section for template 3"""
        skills_header = cell.add_paragraph()
        skills_run = skills_header.add_run("SKILLS")
        skills_run.bold = True
        skills_run.font.size = Pt(12)
        skills_run.font.name = self.font_name
        skills_run.font.color.rgb = self.template3_colors['accent']
        skills_header.paragraph_format.space_after = Pt(8)
        
        # Add underline
        underline_para = cell.add_paragraph()
        underline_run = underline_para.add_run("_" * 30)
        underline_run.font.size = Pt(8)
        underline_run.font.color.rgb = self.template3_colors['line_color']
        underline_para.paragraph_format.space_after = Pt(6)
        
        for skill in self.resume_data.get('skills', []):
            skill_para = cell.add_paragraph()
            skill_name = skill.get('name', '')
            skill_run = skill_para.add_run(f"• {skill_name}")
            skill_run.font.size = Pt(10)
            skill_run.font.name = self.font_name
            skill_para.paragraph_format.space_after = Pt(3)

    def _add_languages_interests_template3(self, cell):
        """Add languages and interests section for template 3"""
        # Languages
        if self.resume_data.get('languages'):
            lang_header = cell.add_paragraph()
            lang_run = lang_header.add_run("LANGUAGES")
            lang_run.bold = True
            lang_run.font.size = Pt(12)
            lang_run.font.name = self.font_name
            lang_run.font.color.rgb = self.template3_colors['accent']
            lang_header.paragraph_format.space_after = Pt(8)
            
            # Add underline
            underline_para = cell.add_paragraph()
            underline_run = underline_para.add_run("_" * 30)
            underline_run.font.size = Pt(8)
            underline_run.font.color.rgb = self.template3_colors['line_color']
            underline_para.paragraph_format.space_after = Pt(6)
            
            for lang in self.resume_data.get('languages', []):
                lang_para = cell.add_paragraph()
                lang_run = lang_para.add_run(f"• {lang}")
                lang_run.font.size = Pt(10)
                lang_run.font.name = self.font_name
                lang_para.paragraph_format.space_after = Pt(3)
        
        # Interests/Hobbies
        interests = self.resume_data.get('interests', []) + self.resume_data.get('hobbies', [])
        if interests:
            int_header = cell.add_paragraph()
            int_run = int_header.add_run("INTERESTS")
            int_run.bold = True
            int_run.font.size = Pt(12)
            int_run.font.name = self.font_name
            int_run.font.color.rgb = self.template3_colors['accent']
            int_header.paragraph_format.space_before = Pt(16)
            int_header.paragraph_format.space_after = Pt(8)
            
            # Add underline
            underline_para = cell.add_paragraph()
            underline_run = underline_para.add_run("_" * 30)
            underline_run.font.size = Pt(8)
            underline_run.font.color.rgb = self.template3_colors['line_color']
            underline_para.paragraph_format.space_after = Pt(6)
            
            for interest in interests:
                int_para = cell.add_paragraph()
                int_run = int_para.add_run(f"• {interest}")
                int_run.font.size = Pt(10)
                int_run.font.name = self.font_name
                int_para.paragraph_format.space_after = Pt(3)

    def _add_circular_image(self, cell, image_data, size=1.5):
        """Add circular profile picture"""
        if not image_data:
            placeholder_para = cell.add_paragraph()
            placeholder_run = placeholder_para.add_run("[Photo]")
            placeholder_run.font.size = Pt(10)
            placeholder_run.font.name = self.font_name
            placeholder_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            return placeholder_para
        
        try:
            # Decode base64 image
            if ',' in image_data:
                image_bytes = base64.b64decode(image_data.split(',')[1])
            else:
                image_bytes = base64.b64decode(image_data)
                
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_img:
                temp_img.write(image_bytes)
                temp_img_path = temp_img.name
            
            # Add to document
            paragraph = cell.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run()
            run.add_picture(temp_img_path, width=Inches(size))
            
            # Cleanup
            os.unlink(temp_img_path)
            return paragraph
        except Exception as e:
            print(f"Error processing profile image: {e}")
            return self._add_circular_image(cell, None, size)